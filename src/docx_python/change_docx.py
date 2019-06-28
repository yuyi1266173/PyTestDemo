# !python3
# -*- coding: utf-8 -*-

from docx import Document


def change_docx_text(old_docx_name, new_docx_name, change_dict):
    old_name_split = old_docx_name.split(".")

    if old_name_split[1] == 'docx':
        document = Document(old_docx_name)

        # tables
        for table in document.tables:
            for row in range(len(table.rows)):
                for col in range(len(table.columns)):
                    for key, value in change_dict.items():
                        if key in table.cell(row, col).text:
                            print(key + "->" + value)
                            table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

        # paragraphs
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                for key, value in change_dict.items():
                    if key in para.runs[i].text:
                        print(key + "->" + value)
                        para.runs[i].text = para.runs[i].text.replace(key, value)

        document.save(new_docx_name)
    else:
        print("不支持的文本格式{}，请转换成docx格式。".format(old_name_split[1]))


def change_report_text():
    pass


if __name__ == "__main__":
    old_name = "儿童行为观察系统报告模板.docx"
    new_name = "儿童行为观察系统报告模板_after.docx"
    change_dict = {
        "张三": "吴宇",
        "2019/05/13": "2019/06/14",
        "替换前的指导建议": "指导建议（替换后）",
        "优秀(替换前)": "非常优秀(替换后)",
    }
    change_docx_text(old_name, new_name, change_dict)

