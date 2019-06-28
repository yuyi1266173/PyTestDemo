# !/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE


def create_demo_docx():
    document = Document()
    document.add_heading('This is my title', 0)
    document.add_paragraph('my paragraph')

    document.styles['Normal'].font.name = u'黑体'
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'我添加的段落文字 ')
    run.font.color.rgb = RGBColor(54, 95, 145)
    run.font.size = Pt(36)

    pic = document.add_picture('rose_logo.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

    rows = 2
    cols = 3
    table = document.add_table(rows=rows, cols=cols, style="Table Grid")  # 添加2行3列的表格

    for i in range(rows):
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "450")
        trPr.append(trHeight)  # 表格高度设置
    # table.autofit = False

    # 设置第二列的宽度
    # col = table.columns[1]
    # col.width = Inches(5)

    arr = [u'序号', u"类型", u"详细描述"]
    heading_cells = table.rows[0].cells
    for i in range(cols):
        p = heading_cells[i].paragraphs[0]
        run = p.add_run(arr[i])
        run.font.color.rgb = RGBColor(54, 95, 145)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 1).text = u'表格文字gerbgtrehbtrhrythv二位跟人谈话人体有合同已经意图具有'
    table.add_row()
    document.save('test1.docx')


def create_docx_demo2():
    # 创建 Document 对象，相当于打开一个 word 文档
    document = Document()

    # 向文档中添加一个标题，标题级别设置为0级
    document.add_heading('This is title', level=0)

    # 向文档中添加一个段落，并将段落引用赋给变量 p
    # 使用 add_run 方法追加字段，并设置格式
    p = document.add_paragraph('This is paragraph')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    # 添加标题和段落，采用不同的形式
    document.add_heading('This is Heading, level 1', level=1)
    document.add_paragraph('Intese quote', style="Intense Quote")
    document.add_paragraph('first item in unordered list', style='List Bullet')
    document.add_paragraph('first item in ordered list', style='List Number')

    # 添加图片，设置图片大小
    document.add_picture(r"rose_logo.png", width=Inches(2.25))

    # 添加表格，填入表格内容
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "cell_00"
    table.cell(0, 1).text = "cell_01"
    table.cell(1, 0).text = "cell_10"
    table.cell(1, 1).text = "cell_11"

    # 保存文本
    document.save('demo.docx')


def list_para_style():
    document = Document()
    styles = document.styles

    # 生成所有段落样式
    for s in styles:
        if s.type == WD_STYLE_TYPE.PARAGRAPH:
            document.add_paragraph('Paragraph style is : ' + s.name, style=s)

    document.save('para_style.docx')


def list_character_style():
    document = Document()
    styles = document.styles
    para = document.add_paragraph()

    # 生成所有字符样式
    for s in styles:
        if s.type == WD_STYLE_TYPE.CHARACTER:
            run = para.add_run("Character style is:  " + s.name + "\n")
            run.style = s

    document.save('character_style.docx')


def list_table_style():
    document = Document()
    styles = document.styles

    # 生成所有表样式
    for s in styles:
        if s.type == WD_STYLE_TYPE.TABLE:
            document.add_paragraph("Table style is :  " + s.name)
            document.add_table(3, 3, style=s)
            document.add_paragraph("\n")

    document.save('table_style.docx')


if __name__ == "__main__":
    # create_demo_docx()
    # create_docx_demo2()
    list_para_style()
    list_character_style()
    list_table_style()
